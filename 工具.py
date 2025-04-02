import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import tkinter as tk
from tkinter import ttk, filedialog

class AwardApp:
    def __init__(self, root):
        self.root = root
        self.root.title("公益学时生成系统")
        self.root.geometry("750x700")
        
        # 初始化变量
        self.input_file = tk.StringVar()
        self.output_path = tk.StringVar()
        self.output_name = tk.StringVar(value="公益学时名单")
        
        # 学时设置
        self.award1 = tk.IntVar(value=4)
        self.award2 = tk.IntVar(value=3)
        self.award3 = tk.IntVar(value=2)
        self.link_staff = tk.BooleanVar(value=True)
        self.staff_award = tk.IntVar(value=self.award2.get())
        
        # 学分设置
        self.credit1 = tk.DoubleVar(value=0.3)
        self.credit2 = tk.DoubleVar(value=0.2)
        self.credit3 = tk.DoubleVar(value=0.1)
        
        # 活动信息
        self.activity_name = tk.StringVar(value="地理科学系·公益活动")
        self.activity_date = tk.StringVar(value="2023-12-31")
        self.department = tk.StringVar(value="学生会公益部")
        self.activity_form = tk.StringVar()
        self.activity_category = tk.StringVar()
        self.guest_count = tk.StringVar()
        self.performer_count = tk.StringVar()
        self.participant_count = tk.StringVar()
        self.audience_count = tk.StringVar()

        # 设置跟踪
        self.award2.trace_add("write", self.update_staff_award)
        self.link_staff.trace_add("write", self.toggle_staff_award)
        
        self.create_widgets()
    
    def update_staff_award(self, *args):
        if self.link_staff.get():
            self.staff_award.set(self.award2.get())
    
    def toggle_staff_award(self, *args):
        if self.link_staff.get():
            self.staff_spinbox.config(state="disabled")
            self.staff_award.set(self.award2.get())
        else:
            self.staff_spinbox.config(state="normal")
    
    def create_widgets(self):
        # 文件选择部分
        ttk.Label(self.root, text="输入文件:").grid(row=0, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(self.root, textvariable=self.input_file, width=35).grid(row=0, column=1, padx=5, pady=2)
        ttk.Button(self.root, text="浏览...", command=self.select_input_file).grid(row=0, column=2, padx=5, pady=2)
        
        # 活动信息部分
        ttk.Label(self.root, text="活动名称 (A):").grid(row=1, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(self.root, textvariable=self.activity_name, width=35).grid(row=1, column=1, padx=5, pady=2)
        
        ttk.Label(self.root, text="活动时间 (B):").grid(row=2, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(self.root, textvariable=self.activity_date, width=35).grid(row=2, column=1, padx=5, pady=2)
        
        ttk.Label(self.root, text="承办部门 (C):").grid(row=3, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(self.root, textvariable=self.department, width=35).grid(row=3, column=1, padx=5, pady=2)
        
        # 其他活动信息
        ttk.Label(self.root, text="活动形式:").grid(row=4, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(self.root, textvariable=self.activity_form, width=35).grid(row=4, column=1, padx=5, pady=2)
        
        ttk.Label(self.root, text="活动类别:").grid(row=5, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(self.root, textvariable=self.activity_category, width=35).grid(row=5, column=1, padx=5, pady=2)
        
        # 学时设置部分
        ttk.Label(self.root, text="一等奖学时:").grid(row=6, column=0, padx=5, pady=2, sticky="e")
        ttk.Spinbox(self.root, from_=1, to=10, textvariable=self.award1, width=5).grid(row=6, column=1, sticky="w")
        ttk.Entry(self.root, textvariable=self.credit1, width=5).grid(row=6, column=2, sticky="w")
        
        ttk.Label(self.root, text="二等奖学时:").grid(row=7, column=0, padx=5, pady=2, sticky="e")
        ttk.Spinbox(self.root, from_=1, to=10, textvariable=self.award2, width=5).grid(row=7, column=1, sticky="w")
        ttk.Entry(self.root, textvariable=self.credit2, width=5).grid(row=7, column=2, sticky="w")
        
        ttk.Label(self.root, text="三等奖学时:").grid(row=8, column=0, padx=5, pady=2, sticky="e")
        ttk.Spinbox(self.root, from_=1, to=10, textvariable=self.award3, width=5).grid(row=8, column=1, sticky="w")
        ttk.Entry(self.root, textvariable=self.credit3, width=5).grid(row=8, column=2, sticky="w")
        
        # 工作人员设置
        ttk.Label(self.root, text="工作人员学时:").grid(row=9, column=0, padx=5, pady=2, sticky="e")
        self.staff_spinbox = ttk.Spinbox(
            self.root, 
            from_=1, 
            to=10, 
            textvariable=self.staff_award, 
            width=5,
            state="disabled"
        )
        self.staff_spinbox.grid(row=9, column=1, sticky="w")
        ttk.Checkbutton(
            self.root, 
            text="与二等奖相同",
            variable=self.link_staff
        ).grid(row=9, column=2, sticky="w", padx=5)
        
        # 其他人数信息
        ttk.Label(self.root, text="特邀嘉宾人数:").grid(row=10, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(self.root, textvariable=self.guest_count, width=35).grid(row=10, column=1, padx=5, pady=2)
        
        ttk.Label(self.root, text="参赛人数:").grid(row=11, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(self.root, textvariable=self.participant_count, width=35).grid(row=11, column=1, padx=5, pady=2)
        
        ttk.Label(self.root, text="观众人数:").grid(row=12, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(self.root, textvariable=self.audience_count, width=35).grid(row=12, column=1, padx=5, pady=2)
        
        # 输出设置
        ttk.Label(self.root, text="保存路径:").grid(row=13, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(self.root, textvariable=self.output_path, width=35).grid(row=13, column=1, padx=5, pady=5)
        ttk.Button(self.root, text="浏览...", command=self.select_output_path).grid(row=13, column=2, padx=5, pady=5)
        
        ttk.Label(self.root, text="文件名称:").grid(row=14, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(self.root, textvariable=self.output_name, width=35).grid(row=14, column=1, padx=5, pady=2)
        
        # 生成按钮
        ttk.Button(self.root, text="生成文档", command=self.generate_document).grid(row=15, column=1, pady=10)
        
        # 状态栏
        self.status = ttk.Label(self.root, text="就绪")
        self.status.grid(row=16, column=0, columnspan=3, sticky="ew")
    
    def select_input_file(self):
        filename = filedialog.askopenfilename(
            filetypes=[("Excel文件", "*.xlsx"), ("所有文件", "*.*")]
        )
        if filename:
            self.input_file.set(filename)
    
    def select_output_path(self):
        path = filedialog.askdirectory()
        if path:
            self.output_path.set(path)
    
    def generate_document(self):
        try:
            if not self.input_file.get():
                raise ValueError("请选择输入文件")
            if not self.output_path.get():
                raise ValueError("请选择保存路径")
            
            filename = f"{self.output_name.get().strip()}.docx"
            full_path = os.path.join(self.output_path.get(), filename)
            os.makedirs(os.path.dirname(full_path), exist_ok=True)
            
            data = process_data(self.input_file.get())
            create_document(
                data=data,
                output_path=full_path,
                awards=[
                    ("一等奖", self.award1.get(), self.credit1.get()),
                    ("二等奖", self.award2.get(), self.credit2.get()),
                    ("三等奖", self.award3.get(), self.credit3.get())
                ],
                staff_award=self.staff_award.get(),
                activity_info={
                    'name': self.activity_name.get(),
                    'date': self.activity_date.get(),
                    'department': self.department.get(),
                    'form': self.activity_form.get(),
                    'category': self.activity_category.get(),
                    'guest': self.guest_count.get(),
                    'performer': self.performer_count.get(),
                    'participant': self.participant_count.get(),
                    'audience': self.audience_count.get()
                }
            )
            
            self.status.config(text=f"文档生成成功：{full_path}", foreground="green")
        except Exception as e:
            self.status.config(text=f"错误：{str(e)}", foreground="red")

def process_data(excel_path):
    df_info = pd.read_excel(excel_path, sheet_name=1, dtype={'十位学号': str})
    df_awards = pd.read_excel(excel_path, sheet_name=0)
    
    merged = pd.merge(
        df_awards[['姓名', '备注（奖项等）']].dropna(),
        df_info[['姓名', '十位学号', '班级']].dropna(),
        on='姓名',
        how='left'
    )
    
    merged['short_id'] = merged['十位学号'].str[-2:].astype(int)
    merged['班级'] = (
        merged['班级'].astype(str).str[:2] + '.' + 
        merged['班级'].astype(str).str[2:].astype(int).astype(str) + '班'
    )
    return merged

def create_document(data, output_path, awards, staff_award, activity_info):
    doc = Document()
    set_global_style(doc)

    # ========== 标题部分 ==========
    title = doc.add_paragraph()
    title_run = title.add_run(f"{activity_info['name']}\n公益学时加分名单")
    title_run.bold = True
    title_run.font.size = Pt(18)  # 小二
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 标题后添加一个空段落
    empty_paragraph = doc.add_paragraph()  # 插入一个空段落
    empty_run = empty_paragraph.add_run()  # 添加一个运行
    empty_run.font.name = '黑体'  # 设置字体为黑体
    empty_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 设置中文字体为黑体
    empty_run.font.size = Pt(18)  # 设置字体大小为小二（18磅）

    # ========== 表格部分 ==========
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'

    # 设置列宽（3cm,4cm,3cm,4.5cm）
    widths = (3, 4, 3, 4.5)
    for i, width in enumerate(widths):
        table.columns[i].width = int(width * 567)  # 厘米转缇

    # 填充表格数据
    cells = [
        ("活动日期", activity_info['date'], "活动地点", ""),
        ("活动形式", activity_info['form'], "活动类别", activity_info['category']),
        ("承办部门", activity_info['department'], "工作人员人数", str(len(data[data['备注（奖项等）'] == '工作人员']))),
        ("特邀嘉宾人数", activity_info['guest'], "特邀表演人数", activity_info['performer']),
        ("活动获奖人数", str(sum(len(data[data['备注（奖项等）'].str.contains(level)]) for level, _, _ in awards)), "参赛（参与）人数", activity_info['participant']),
        ("活动观众人数", activity_info['audience'], "加学时总人数", str(len(data))),
        ("申请学时梯度", "/".join(str(award[1]) for award in awards), "申请学分梯度", "/".join(f"{award[2]:.1f}" for award in awards))
    ]

    for row_idx, row_data in enumerate(cells):
        row = table.rows[row_idx]
        row.height = int(0.85 * 567)  # 行高0.85cm
        
        for col_idx, cell_data in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_data)
            cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if col_idx % 2 == 0:
                paragraph.runs[0].bold = True

    # 表格后添加小四回车（12磅间距）
    mpty_paragraph = doc.add_paragraph()  # 插入一个空段落
    empty_run = empty_paragraph.add_run()  # 添加一个运行
    empty_run.font.name = '黑体'  # 设置字体为黑体
    empty_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 设置中文字体为黑体
    empty_run.font.size = Pt(12)  # 设置字体大小为小二（12磅）

    # ========== 正文内容 ==========
    # 处理工作人员
    staff = data[data['备注（奖项等）'] == '工作人员']
    if not staff.empty:
        total_staff = add_section(doc, "一、工作人员加分：", staff, staff_award)
        add_remark(doc, total_staff, staff_award)
    
    # 处理获奖人员
    awards_data = data[data['备注（奖项等）'].str.contains('等奖')]
    if not awards_data.empty:
        add_awards_section(doc, awards_data, awards)
    
    doc.save(output_path)

def set_global_style(doc):
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_section(doc, title, data, credit):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    p.paragraph_format.space_after = Pt(6)
    
    total = 0
    for (cls, group) in data.groupby('班级', sort=False):
        students = group.sort_values('short_id').apply(format_student, axis=1).tolist()
        total += len(students)
        chunks = [students[i:i+7] for i in range(0, len(students), 7)]
        
        for i, chunk in enumerate(chunks):
            para = doc.add_paragraph()
            
            if i == 0:
                cls_run = para.add_run(cls)
                add_space(para, count=2)
            else:
                add_space(para, count=6, font_size=9)
                add_space(para, count=4)
            
            for item in chunk:
                para.add_run(item['formatted_name'])
                para.add_run(str(item['id'])).font.size = Pt(9)
                para.add_run(item['space']).font.size = Pt(9)
    
    return total

def add_awards_section(doc, data, awards):
    p = doc.add_paragraph()
    run = p.add_run("二、获奖人员加分：")
    run.bold = True
    p.paragraph_format.space_after = Pt(6)
    
    for level, credit, _ in awards:
        level_data = data[data['备注（奖项等）'].str.contains(level)]
        if not level_data.empty:
            total = add_section(doc, f"{level}人员加分：", level_data, credit)
            add_remark(doc, total, credit)

def add_remark(doc, count, credit):
    remark = f"（以上{count}名同学各加{credit}公益学时）"
    p = doc.add_paragraph(remark)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def add_space(para, count, font_size=12):
    space_run = para.add_run(' ' * count)
    space_run.font.size = Pt(font_size)

def format_student(row):
    id_num = row['short_id']
    name = row['姓名']
    
    if len(name) == 2:
        formatted_name = f"{name[0]}  {name[1]}"
    else:
        formatted_name = name
    
    return {
        'formatted_name': formatted_name,
        'id': str(id_num),
        'space': '  ' if id_num < 10 else ' '
    }

if __name__ == "__main__":
    root = tk.Tk()
    app = AwardApp(root)
    root.mainloop()